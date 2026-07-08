#!/usr/bin/env python3
"""
将 Markdown 文档转换为格式化的 Word 文档
支持标题、列表、表格、代码块等
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def setup_styles(doc):
    """设置文档样式（黑色字体，微软雅黑）"""
    # 正文样式
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(11)
    font.color.rgb = RGBColor(0, 0, 0)

    # 设置中文字体
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # 标题样式
    heading_sizes = {1: 18, 2: 16, 3: 14, 4: 12}
    for level, size in heading_sizes.items():
        try:
            heading_style = doc.styles[f'Heading {level}']
            heading_font = heading_style.font
            heading_font.name = '微软雅黑'
            heading_font.size = Pt(size)
            heading_font.color.rgb = RGBColor(0, 0, 0)
            heading_font.bold = True
            heading_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        except KeyError:
            pass


def add_paragraph_with_style(doc, text, style='Normal'):
    """添加带样式的段落"""
    para = doc.add_paragraph(text, style=style)
    for run in para.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
    return para


def parse_table(lines):
    """解析 Markdown 表格"""
    rows = []
    for line in lines:
        line = line.strip()
        if not line or line.startswith('|---') or line.startswith('| ---'):
            continue
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        rows.append(cells)
    return rows


def add_table_to_doc(doc, rows):
    """添加表格到文档"""
    if not rows:
        return

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = 'Light Grid Accent 1'

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            cell = table.rows[i].cells[j]
            cell.text = cell_text
            # 设置单元格字体
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                    run.font.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                # 表头加粗
                if i == 0:
                    for run in paragraph.runs:
                        run.font.bold = True


def convert_md_to_docx(md_file, docx_file):
    """转换 Markdown 到 Word"""
    print(f'Converting {md_file} to {docx_file}...')

    with open(md_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    doc = Document()
    setup_styles(doc)

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_lines = []

    while i < len(lines):
        line = lines[i].rstrip()

        # 代码块
        if line.startswith('```'):
            if in_code_block:
                # 结束代码块
                code_text = '\n'.join(code_lines)
                para = doc.add_paragraph(code_text)
                para.style = 'Normal'
                for run in para.runs:
                    run.font.name = 'Consolas'
                    run.font.size = Pt(9)
                    run.font.color.rgb = RGBColor(0, 0, 0)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # 表格检测
        if line.startswith('|') and '|' in line[1:]:
            if not in_table:
                in_table = True
                table_lines = []
            table_lines.append(line)
            i += 1
            continue
        elif in_table and not line.startswith('|'):
            # 表格结束
            rows = parse_table(table_lines)
            add_table_to_doc(doc, rows)
            doc.add_paragraph()  # 表格后加空行
            in_table = False
            table_lines = []
            # 不要 continue，处理当前行

        # 标题
        if line.startswith('#'):
            level = len(line) - len(line.lstrip('#'))
            text = line.lstrip('#').strip()
            if level <= 4:
                add_paragraph_with_style(doc, text, f'Heading {level}')
            else:
                add_paragraph_with_style(doc, text, 'Heading 4')
        # 无序列表
        elif line.startswith('- ') or line.startswith('* '):
            text = line[2:].strip()
            para = doc.add_paragraph(text, style='List Bullet')
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        # 有序列表
        elif re.match(r'^\d+\.\s', line):
            text = re.sub(r'^\d+\.\s', '', line).strip()
            para = doc.add_paragraph(text, style='List Number')
            for run in para.runs:
                run.font.color.rgb = RGBColor(0, 0, 0)
        # 空行
        elif not line:
            doc.add_paragraph()
        # 普通段落
        else:
            # 处理行内代码 `code`
            text = re.sub(r'`([^`]+)`', r'\1', line)
            # 处理粗体 **text**
            text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
            add_paragraph_with_style(doc, text)

        i += 1

    # 处理未结束的表格
    if in_table:
        rows = parse_table(table_lines)
        add_table_to_doc(doc, rows)

    doc.save(docx_file)
    print(f'✓ Saved to {docx_file}')


if __name__ == '__main__':
    import sys

    files_to_convert = [
        ('/Users/zeal/Desktop/ELM/docs/需求说明书.md', '/Users/zeal/Desktop/ELM/docs/需求说明书.docx'),
        ('/Users/zeal/Desktop/ELM/docs/设计说明书.md', '/Users/zeal/Desktop/ELM/docs/设计说明书.docx'),
    ]

    for md_file, docx_file in files_to_convert:
        try:
            convert_md_to_docx(md_file, docx_file)
        except Exception as e:
            print(f'✗ Error converting {md_file}: {e}')
            import traceback
            traceback.print_exc()
